"""Read plan text from uploaded txt/docx/pdf files."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

import pdfplumber
from docx import Document


def _decode_text(data: bytes) -> str:
    """Decode text files with common Chinese encodings."""
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _read_docx(data: bytes) -> str:
    try:
        document = Document(BytesIO(data))
    except Exception:
        return _read_docx_xml_fallback(data)

    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    return "\n".join(parts)


def _read_docx_xml_fallback(data: bytes) -> str:
    """Extract text from malformed docx files without following broken media rels."""
    with ZipFile(BytesIO(data)) as archive:
        document_xml = archive.read("word/document.xml")

    root = ET.fromstring(document_xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    parts: list[str] = []
    for paragraph in root.findall(".//w:p", ns):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns)).strip()
        if text:
            parts.append(text)
    return "\n".join(parts)


def _read_pdf(data: bytes) -> str:
    parts: list[str] = []
    with pdfplumber.open(BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text.strip())
    return "\n\n".join(parts)


def read_text_from_upload(uploaded_file) -> str:
    """Read text from a Streamlit uploaded file."""
    filename = uploaded_file.name.lower()
    data = uploaded_file.getvalue()

    if filename.endswith(".txt"):
        return _decode_text(data)
    if filename.endswith(".docx"):
        return _read_docx(data)
    if filename.endswith(".pdf"):
        return _read_pdf(data)

    raise ValueError("仅支持 .txt / .docx / .pdf 文件")


def read_text_from_path(path: str | Path) -> str:
    """Read text from a local path. Useful for manual checks."""
    path = Path(path)
    data = path.read_bytes()
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return _decode_text(data)
    if suffix == ".docx":
        return _read_docx(data)
    if suffix == ".pdf":
        return _read_pdf(data)

    raise ValueError("仅支持 .txt / .docx / .pdf 文件")
